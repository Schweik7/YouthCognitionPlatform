#!/usr/bin/env python3
"""
解析 Word 文档，提取题目和图片
"""

import os
import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches
import io
from PIL import Image


def extract_images_from_docx(doc_path, output_dir):
    """从 Word 文档中提取图片"""
    doc = Document(doc_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    image_files = []
    image_counter = 1
    
    # 遍历文档中的所有关系
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                # 获取图片数据
                image_data = rel.target_part.blob
                
                # 确定图片格式
                if image_data.startswith(b'\xff\xd8'):
                    ext = '.jpg'
                elif image_data.startswith(b'\x89PNG'):
                    ext = '.png'
                elif image_data.startswith(b'GIF'):
                    ext = '.gif'
                else:
                    ext = '.png'  # 默认为 png
                
                # 保存图片
                image_filename = f"question{image_counter}{ext}"
                image_path = output_dir / image_filename
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                image_files.append(f"images/reading/{image_filename}")
                image_counter += 1
                print(f"提取图片: {image_filename}")
                
            except Exception as e:
                print(f"提取图片时出错: {e}")
                continue
    
    return image_files


def extract_questions_from_table(doc_path):
    """从 Word 文档的第二个表格中提取题目"""
    doc = Document(doc_path)
    
    if len(doc.tables) < 2:
        print("文档中没有找到第二个表格")
        return []
    
    # 获取第二个表格
    table = doc.tables[1]
    questions = []
    
    # 跳过表头，从第二行开始
    for i, row in enumerate(table.rows[1:], 1):
        if len(row.cells) >= 2:
            # 第一列是题目序号，第二列是题目内容
            try:
                question_id = row.cells[0].text.strip()
                question_text = row.cells[1].text.strip()
                
                if question_id and question_text:
                    # 检查题目中是否包含图片（通过检查是否有嵌入的图片或特定关键词）
                    has_image = any(keyword in question_text for keyword in ['图', '图片', '图中', '从图', '照片', '画面'])
                    
                    questions.append({
                        'id': int(question_id) if question_id.isdigit() else i,
                        'text': question_text,
                        'has_image': has_image
                    })
                    
            except Exception as e:
                print(f"解析第 {i} 行时出错: {e}")
                continue
    
    return questions


def generate_text_answers(questions):
    """为纯文字题目生成答案（基于常识判断）"""
    answers = {}
    
    for q in questions:
        if not q['has_image']:
            text = q['text'].lower()
            question_id = q['id']
            
            # 基于常识的简单判断规则
            if any(keyword in text for keyword in ['太阳从西边升起', '月亮发光', '地球是平的', '冰在零度以上', '火车比汽车慢']):
                answers[question_id] = False
            elif any(keyword in text for keyword in ['植物光合作用', '声音在真空', '磁铁同极相吸', '电流从正极', '光在水中']):
                answers[question_id] = False
            elif any(keyword in text for keyword in ['海底住企鹅', '鲸鱼最小', '昆虫八条腿', '鸟类不游泳', '液体不结冰', '重力向上']):
                answers[question_id] = False
            elif any(keyword in text for keyword in ['小猫在鱼缸游泳', '雪花是热的']):
                answers[question_id] = False
            else:
                # 默认为正确（大部分描述性陈述）
                answers[question_id] = True
    
    return answers


def create_csv_file(questions, image_files, text_answers, csv_path):
    """创建 CSV 文件"""
    csv_path = Path(csv_path)
    
    with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['id', 'text', 'image_path', 'correct_answer'])
        
        image_index = 0
        
        for q in questions:
            question_id = q['id']
            text = q['text']
            
            # 确定图片路径
            image_path = ''
            if q['has_image'] and image_index < len(image_files):
                image_path = image_files[image_index]
                image_index += 1
            
            # 确定答案
            if q['has_image']:
                # 有图片的题目，答案待用户提供
                correct_answer = 'USER_TO_PROVIDE'
            else:
                # 纯文字题目，使用生成的答案
                correct_answer = 'TRUE' if text_answers.get(question_id, True) else 'FALSE'
            
            writer.writerow([question_id, text, image_path, correct_answer])


def main():
    """主函数"""
    # 文件路径
    doc_path = Path('apps/reading_fluency/3min阅读图画版.docx')
    image_output_dir = Path('../public/images/reading')
    csv_output_path = Path('data/初中及以上阅读.csv')
    
    if not doc_path.exists():
        print(f"找不到文档文件: {doc_path}")
        return
    
    print("开始解析 Word 文档...")
    
    # 1. 提取图片
    print("提取图片...")
    image_files = extract_images_from_docx(doc_path, image_output_dir)
    print(f"提取了 {len(image_files)} 张图片")
    
    # 2. 提取题目
    print("提取题目...")
    questions = extract_questions_from_table(doc_path)
    print(f"提取了 {len(questions)} 道题目")
    
    # 3. 为纯文字题目生成答案
    print("生成文字题目答案...")
    text_answers = generate_text_answers(questions)
    print(f"为 {len(text_answers)} 道纯文字题目生成了答案")
    
    # 4. 创建 CSV 文件
    print("创建 CSV 文件...")
    create_csv_file(questions, image_files, text_answers, csv_output_path)
    print(f"CSV 文件已保存到: {csv_output_path}")
    
    # 5. 显示统计信息
    image_questions = sum(1 for q in questions if q['has_image'])
    text_questions = len(questions) - image_questions
    
    print(f"\n统计信息:")
    print(f"总题目数: {len(questions)}")
    print(f"带图片题目: {image_questions}")
    print(f"纯文字题目: {text_questions}")
    print(f"提取图片数: {len(image_files)}")
    
    print(f"\n注意: 带图片的题目答案标记为 'USER_TO_PROVIDE'，请手动更新这些答案。")


if __name__ == '__main__':
    main()